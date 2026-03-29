import docx
import sys

doc = docx.Document(r'c:\Users\DELL\Desktop\قانون\القانون المدني المصري .docx')
lines = []

for p in doc.paragraphs:
    text = p.text.strip()
    if text:
        lines.append(text)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            text = cell.text.strip()
            if text:
                lines.append(text)

with open(r'c:\Users\DELL\Desktop\قانون\out_sys.txt', 'w', encoding='utf-8') as f:
    f.write('\n'.join(lines[:200]))
